# import os
# import re
# import json
# import spacy
# import pdfplumber
# import docx
# import pandas as pd
# from pathlib import Path
# from tkinter import filedialog, Tk, Label, Button, messagebox

# nlp = spacy.load("en_core_web_sm")

# def extract_text_from_pdf(file_path):
#     text = ""
#     with pdfplumber.open(file_path) as pdf:
#         for page in pdf.pages:
#             text += page.extract_text() or ""
#     return text

# def extract_text_from_docx(file_path):
#     doc = docx.Document(file_path)
#     return "\n".join([para.text for para in doc.paragraphs])

# def extract_email(text):
#     match = re.findall(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', text)
#     return match[0] if match else ""

# def extract_phone(text):
#     match = re.findall(r'\+?\d[\d\s()-]{8,}\d', text)
#     return match[0] if match else ""

# def extract_entities(text):
#     doc = nlp(text)
#     name = ""
#     experience = ""
#     skills = []

#     # Step 1: Extract Name from first few lines
#     lines = text.strip().split('\n')
#     for line in lines:
#         if len(line.strip().split()) == 2:
#             name = line.strip()
#             break

#     # Step 2: Regex based Experience
#     exp_match = re.search(r'(\d+)\s+(years|year)\s+(of|in)?\s?(experience|exp|working)?\s?(at|in|on)?\s?(.*)', text, re.IGNORECASE)
#     if exp_match:
#         experience = exp_match.group(0)

#     # Step 3: Skill Matching from Predefined List
#     skill_keywords = [
#         "Python", "Java", "JavaScript", "C++", "SQL", "Django", "Flask",
#         "HTML", "CSS", "Excel", "Power BI", "Pandas", "NumPy", "Git"
#     ]
#     for skill in skill_keywords:
#         if re.search(rf'\b{re.escape(skill)}\b', text, re.IGNORECASE):
#             skills.append(skill)

#     return name.strip(), list(set(skills)), experience.strip()


# def parse_resume(file_path):
#     ext = Path(file_path).suffix
#     text = ""
#     if ext == ".pdf":
#         text = extract_text_from_pdf(file_path)
#     elif ext == ".docx":
#         text = extract_text_from_docx(file_path)
#     else:
#         return None

#     email = extract_email(text)
#     phone = extract_phone(text)
#     name, skills, experience = extract_entities(text)

#     return {
#         "File": os.path.basename(file_path),
#         "Name": name,
#         "Email": email,
#         "Phone": phone,
#         "Experience": experience,
#         "Skills": skills
#     }

# def process_files(file_paths):
#     all_data = []
#     for file_path in file_paths:
#         data = parse_resume(file_path)
#         if data:
#             all_data.append(data)

#     os.makedirs("output", exist_ok=True)

#     # Save JSON
#     with open("output/parsed_resumes.json", "w") as f:
#         json.dump(all_data, f, indent=4)

#     # Save Excel
#     df = pd.DataFrame(all_data)
#     df.to_excel("output/parsed_resumes.xlsx", index=False)

#     messagebox.showinfo("Done", "✅ Resumes parsed and saved to 'output/' folder.")

# def browse_files():
#     file_paths = filedialog.askopenfilenames(filetypes=[("Resumes", "*.pdf *.docx")])
#     if file_paths:
#         process_files(file_paths)

# # GUI
# app = Tk()
# app.title("AI Resume Parser")
# app.geometry("400x200")
# app.resizable(False, False)

# Label(app, text="AI-Based Resume Parser", font=("Helvetica", 16, "bold")).pack(pady=20)
# Button(app, text="Select Resumes (PDF/DOCX)", command=browse_files, width=30).pack(pady=10)
# Button(app, text="Exit", command=app.quit, width=15).pack(pady=10)

# app.mainloop()






import os
import re
import json
import spacy
import pdfplumber
import docx
import pandas as pd
from pathlib import Path
from tkinter import filedialog, Tk, Label, Button, messagebox

nlp = spacy.load("en_core_web_sm")

# Extract from PDF
def extract_text_from_pdf(file_path):
    text = ""
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text += page.extract_text() or ""
    return text

# Extract from DOCX
def extract_text_from_docx(file_path):
    doc = docx.Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])

# Email extraction
def extract_email(text):
    match = re.findall(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', text)
    return match[0] if match else ""

# Phone extraction
def extract_phone(text):
    match = re.findall(r'\+?\d[\d\s()-]{8,}\d', text)
    return match[0] if match else ""

# Entity extraction (name, skills, experience)
def extract_entities(text):
    doc = nlp(text)
    name = ""
    experience = ""
    skills = []

    # Step 1: Name from top line
    lines = text.strip().split('\n')
    for line in lines:
        if len(line.strip().split()) == 2:
            name = line.strip()
            break

    # Step 2: Experience via regex
    exp_match = re.search(r'(\d+)\s+(years|year)\s+(of|in)?\s?(experience|exp|working)?\s?(at|in|on)?\s?(.*)', text, re.IGNORECASE)
    if exp_match:
        experience = exp_match.group(0)

    # Step 3: Skill Matching (expandable)
    skill_keywords = [
        "Python", "Java", "JavaScript", "C++", "SQL", "Django", "Flask",
        "HTML", "CSS", "Excel", "Power BI", "Pandas", "NumPy", "Git"
    ]
    for skill in skill_keywords:
        if re.search(rf'\b{re.escape(skill)}\b', text, re.IGNORECASE):
            skills.append(skill)

    return name.strip(), list(set(skills)), experience.strip()

# Parse single resume
def parse_resume(file_path):
    ext = Path(file_path).suffix.lower()
    text = ""
    if ext == ".pdf":
        text = extract_text_from_pdf(file_path)
    elif ext == ".docx":
        text = extract_text_from_docx(file_path)
    else:
        return None

    email = extract_email(text)
    phone = extract_phone(text)
    name, skills, experience = extract_entities(text)

    return {
        "File": os.path.basename(file_path),
        "Name": name,
        "Email": email,
        "Phone": phone,
        "Experience": experience,
        "Skills": skills
    }

# Process all resumes
def process_files(file_paths):
    all_data = []
    for file_path in file_paths:
        data = parse_resume(file_path)
        if data:
            all_data.append(data)

    os.makedirs("output", exist_ok=True)

    # Remove old files if exist
    if os.path.exists("output/parsed_resumes.json"):
        os.remove("output/parsed_resumes.json")
    if os.path.exists("output/parsed_resumes.xlsx"):
        os.remove("output/parsed_resumes.xlsx")

    # Save as JSON
    with open("output/parsed_resumes.json", "w", encoding='utf-8') as f:
        json.dump(all_data, f, indent=4)

    # Save as Excel
    df = pd.DataFrame(all_data)
    df["Skills"] = df["Skills"].apply(lambda x: ', '.join(x))  # Clean skill list
    df.to_excel("output/parsed_resumes.xlsx", index=False)

    messagebox.showinfo("Done", f"✅ Parsed {len(all_data)} resumes.\nSaved to 'output/' folder.")

# Browse full folder (batch parse)
def browse_folder():
    folder_path = filedialog.askdirectory(title="Select Resume Folder")
    if folder_path:
        file_paths = []
        for ext in ("*.pdf", "*.docx"):
            file_paths.extend(Path(folder_path).glob(ext))
        file_paths = [str(fp) for fp in file_paths]
        if file_paths:
            process_files(file_paths)
        else:
            messagebox.showwarning("No Resumes", "⚠️ No PDF or DOCX files found in the folder.")

# GUI...........................................
app = Tk()
app.title("AI Resume Parser")
app.geometry("420x220")
app.resizable(False, False)

Label(app, text="AI-Based Resume Parser", font=("Helvetica", 16, "bold")).pack(pady=20)
Button(app, text="Select Resume Folder (PDF & DOCX)", command=browse_folder, width=40).pack(pady=10)
Button(app, text="Exit", command=app.quit, width=20).pack(pady=10)

app.mainloop()
